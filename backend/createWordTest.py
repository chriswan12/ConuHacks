from docx import Document
from docx.text.paragraph import Paragraph

document = Document()
document.add_heading("this is the title", 0)


p = document.add_paragraph('This is the body')
p.add_run('some bold text').bold = True
p.add_run('and italic text.').italic = True
document.save('output.docx')

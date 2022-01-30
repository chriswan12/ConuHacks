from docx import Document
import datetime
from docx.text.paragraph import Paragraph
import smtplib
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email import encoders


class GenerateMSWord:
    def __init__(self, text, title=None):
        self.text = text
        self.title = title

    def generate_docx(self, email):
        document = Document()
        if self.title:
            document.add_heading(self.title, 1)
        document.add_paragraph(str(datetime.datetime.now().date()))
        document.add_paragraph(self.text)

        document.save('summary.docx')
        self._send_email(email, files=['summary.docx'])
        return

    def _send_email(self, receiver, email='conuhacks2022@gmail.com',
                    passwd='ConUHacks2022', files=None):  # files = []
        if not receiver:
            return False
        msg = MIMEMultipart()
        msg['From'] = email
        msg['To'] = receiver
        msg['Subject'] = 'Summary'
        body = 'Below the summary is attached.'
        msg.attach(MIMEText(body))

        if files:
            for filename in files:

                attachment = open(filename, 'rb')

                part = MIMEBase("application", "octet-stream")

                part.set_payload(attachment.read())

                encoders.encode_base64(part)

                part.add_header("Content-Disposition",
                                f"attachment; filename= {filename}")

                msg.attach(part)

        msg = msg.as_string()

        try:
            server = smtplib.SMTP('smtp.gmail.com:587')
            server.ehlo()
            server.starttls()
            server.login(email, passwd)
            server.sendmail(email, receiver, msg)
            server.quit()
            print('Email sent successfully')

        except:
            print("Email couldn't be sent")


if __name__ == "__main__":
    with open('../formatted.txt', 'r') as f:
        text = f.readline()
        word = GenerateMSWord(text,
                              "Test Summary of Part of the Lecture").generate_docx(
            'rekiyo4126@peykesabz.com')

